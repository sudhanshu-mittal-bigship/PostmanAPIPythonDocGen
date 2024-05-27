import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_word_document(json_file_path, output_file_path):
    def get_unique_filename(file_path):
        filename, extension = os.path.splitext(file_path)
        count = 1
        while os.path.exists(file_path):
            file_path = f"{filename}_{count}{extension}"
            count += 1
        return file_path

    output_file_path = get_unique_filename(output_file_path)

    # Load the JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)

    # Create a new Document
    doc = Document()

    # Add the main title and description
    info = data.get('info', {})
    main_title = info.get('name', 'API Documentation')
    main_description = info.get('description', '')

    # Add a title
    doc.add_heading('API Documentation', 0)
    doc.add_heading(main_title, 0)

    if main_description:
        doc.add_paragraph(main_description)

    # Function to add sections
    def add_section(title, content, level=1):
        heading = doc.add_heading(title, level=level)
        p = doc.add_paragraph(content)
        p.style.font.size = Pt(12)

    # Function to format JSON payloads
    def format_json_payload(json_payload, level=1):
        def add_colored_text(paragraph, text, color):
            run = paragraph.add_run(text)
            run.font.color.rgb = color
            run.font.size = Pt(12)

        def process_json(json_obj, paragraph):
            if isinstance(json_obj, dict):
                for key, value in json_obj.items():
                    add_colored_text(paragraph, f'"{key}": ', RGBColor(255, 0, 0))  # Red for keys
                    if isinstance(value, (dict, list)):
                        paragraph.add_run('\n')
                        process_json(value, paragraph)
                    else:
                        add_colored_text(paragraph, f'"{value}",\n', RGBColor(0, 128, 0))  # Green for values
            elif isinstance(json_obj, list):
                for item in json_obj:
                    process_json(item, paragraph)
            else:
                add_colored_text(paragraph, f'"{json_obj}",\n', RGBColor(0, 128, 0))  # Green for values

        json_paragraph = doc.add_paragraph()
        process_json(json_payload, json_paragraph)
        json_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Recursive function to process items
    def process_items(items, level=1):
        for item in items:
            request_name = item.get('name', 'Unnamed Request')
            add_section(request_name, '', level)

            if 'request' in item:
                request = item['request']
                request_description = request.get('description', 'No description')
                method = request.get('method', 'GET')
                url = request.get('url', {}).get('raw', 'No URL')
                headers = request.get('header', [])
                body = request.get('body', {}).get('raw', '')

                add_section('Description', request_description, level + 1)
                add_section('Method', method, level + 1)
                add_section('URL', url, level + 1)
                
                # Add headers
                headers_text = "\n".join([f"{header['key']}: {header['value']}" for header in headers])
                add_section('Headers', headers_text, level + 1)

                # Add body
                if body:
                    try:
                        json_body = json.loads(body)
                        add_section('Body', '', level + 1)
                        format_json_payload(json_body, level + 2)
                    except json.JSONDecodeError:
                        add_section('Body', body, level + 1)

            if 'item' in item:
                process_items(item['item'], level + 1)

    # Process the top-level items
    process_items(data['item'], level=1)

    # Save the document
    doc.save(output_file_path)
    print(f"Document saved to {output_file_path}")

# Replace with your file paths
json_file_path = 'C:/xampp/htdocs/python/exported_collection.json'
output_file_path = 'C:/Users/sudha/Downloads/outbound_api_document.docx'

# Create the Word document
create_word_document(json_file_path, output_file_path)
